from flask import Flask, render_template, request, send_from_directory
from docx import Document
import os

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Word dosyası oluşturma işlemleri
        doc = Document()
        doc.add_heading('Test Raporu', 0)

        doc.add_heading('Test Senaryosu:', level=1)
        doc.add_paragraph(request.form.get('scenario'))

        doc.add_heading('1. Test Senaryosu Adı:', level=1)
        doc.add_paragraph(request.form.get('scenario_name'))

        doc.add_heading('Adımlar:', level=1)
        doc.add_paragraph(request.form.get('steps'))
        for step in request.form.getlist('extraSteps'):
            doc.add_paragraph(step, style='BodyText')

        doc.add_heading('Beklenen Sonuç:', level=1)
        doc.add_paragraph(request.form.get('expected_result'))
        for exp in request.form.getlist('extraExpectedResult'):
            doc.add_paragraph(exp, style='BodyText')

        doc.add_heading('Gerçek Sonuç:', level=1)
        doc.add_paragraph(request.form.get('actual_result'))
        for act in request.form.getlist('extraActualResult'):
            doc.add_paragraph(act, style='BodyText')

        doc.add_heading('Durum:', level=1)
        doc.add_paragraph(request.form.get('status'))

        filename = "rapor.docx"
        doc.save(filename)
        return send_from_directory(os.getcwd(), filename, as_attachment=True)

    return render_template('index.html')  # Bu kısım, yukarıda verdiğiniz HTML dosyasını 'index.html' adıyla templates klasöründe saklamalısınız.

if __name__ == "__main__":
    app.run(debug=True)
